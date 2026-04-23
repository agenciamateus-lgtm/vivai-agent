"""
generate_docx.py — Proposta Comercial VIVAI Studio
Modelo fiel ao PDF original: logo, tabela, caixa objeto, títulos cyan, footer.
"""
import sys, json, os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENTATION
from datetime import datetime

# ── Cores ──────────────────────────────────────────────────────────────────
CYAN  = RGBColor(0x00, 0x90, 0xBB)   # azul-cyan dos títulos
NAVY  = RGBColor(0x1A, 0x1A, 0x2E)   # quase preto
GRAY  = RGBColor(0x55, 0x55, 0x55)   # texto secundário
LGRAY = RGBColor(0xAA, 0xAA, 0xAA)   # linhas
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def brl(v):
    try:
        return f"R$ {float(v or 0):,.2f}".replace(",","X").replace(".",",").replace("X",".")
    except:
        return "R$ 0,00"

# ── Helpers XML ────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6.upper())
    tcPr.append(shd)

def set_cell_borders(cell, sides='all', color='CCCCCC', size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    all_sides = ['top','left','bottom','right'] if sides == 'all' else sides
    for side in all_sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color.upper())
        tcBorders.append(el)
    # Set inner borders to none if only outer
    if sides != 'all':
        for side in ['top','left','bottom','right']:
            if side not in all_sides:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                tcBorders.append(el)
    tcPr.append(tcBorders)

def cell_padding(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)

def add_run(para, text, bold=False, italic=False, color=BLACK, size=10.5, font='Calibri'):
    run = para.add_run(text)
    run.font.name  = font
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run

def new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def section_title(doc, number, title):
    """Título de seção: '1. ESCOPO DOS SERVIÇOS' em cyan bold"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(5)
    add_run(p, f"{number}. {title.upper()}", bold=True, color=CYAN, size=11, font='Calibri')
    return p

def body_para(doc, text, space_before=0, space_after=5, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    add_run(p, text, color=BLACK, size=10.5)
    return p

def bullet_item(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    add_run(p, "• ", bold=False, color=CYAN, size=10.5)
    add_run(p, text, color=BLACK, size=10.5)
    return p

# ── Gerador principal ──────────────────────────────────────────────────────
def generate(data, output_path, logo_path=None):
    doc = Document()

    # Margens (2.5cm laterais, 2cm top/bottom)
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)
        sec.page_width    = Cm(21.0)   # A4
        sec.page_height   = Cm(29.7)

    # Estilo padrão
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)

    # Dados
    client   = data.get('client','')
    event    = data.get('event','')
    date     = data.get('date','')
    location = data.get('location','')
    cnpj     = data.get('cnpj','')
    contact  = data.get('contact','')
    delivery = data.get('delivery','')
    notes    = data.get('notes','')
    total    = data.get('total',0)
    subtotal = data.get('subtotal',0)
    caixa    = data.get('caixa',0)
    imposto  = data.get('imposto',0)
    transp   = data.get('transport_total',0)
    edicao_t = data.get('edicao_total',0)
    diaria_t = data.get('diaria_team',0)
    diarias  = int(data.get('diarias',1))
    team_items   = data.get('team_items',[])
    edicao_items = data.get('edicao_items',[])
    parcelas     = data.get('parcelas',[])
    escopo       = data.get('escopo',[])
    entregaveis  = data.get('entregaveis',[])
    objeto = data.get('objeto') or f"Prestação de serviços audiovisuais para {event or 'produção audiovisual'} da {client}, conforme escopo e condições descritos nesta proposta."

    # ── FOOTER em todas as páginas ────────────────────────────────────────
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(fp, f"VIVAI STUDIO  •  Proposta Comercial  •  {client or 'Cliente'}", 
            color=GRAY, size=8.5)
    # Linha acima do footer
    fpPr = fp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'),'single'); top.set(qn('w:sz'),'4')
    top.set(qn('w:space'),'3'); top.set(qn('w:color'),'CCCCCC')
    pBdr.append(top); fpPr.append(pBdr)

    # ── CABEÇALHO: logo + título ──────────────────────────────────────────
    # Tabela 2 colunas: logo | título
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.style = 'Table Grid'
    # Remove all borders
    for row in hdr_tbl.rows:
        for cell in row.cells:
            set_cell_borders(cell, sides=[], color='FFFFFF', size=0)
            # Actually set all borders to none
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(old)
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right','insideH','insideV']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                tcBorders.append(el)
            tcPr.append(tcBorders)

    hdr_tbl.columns[0].width = Cm(4)
    hdr_tbl.columns[1].width = Cm(12.5)

    # Logo cell
    logo_cell = hdr_tbl.rows[0].cells[0]
    logo_cell._tc.get_or_add_tcPr()
    lp = logo_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_after = Pt(0)

    actual_logo = logo_path
    if not actual_logo or not os.path.exists(actual_logo):
        for candidate in ['/app/logo.png','logo.png','/mnt/user-data/outputs/logo.png']:
            if os.path.exists(candidate):
                actual_logo = candidate
                break

    if actual_logo and os.path.exists(actual_logo):
        from docx.shared import Inches
        run_logo = lp.add_run()
        run_logo.add_picture(actual_logo, width=Cm(3.2))

    # Title cell
    title_cell = hdr_tbl.rows[0].cells[1]
    title_cell._tc.get_or_add_tcPr()
    # vertical align middle
    tcPr = title_cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

    tp1 = title_cell.paragraphs[0]
    tp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp1.paragraph_format.space_after = Pt(4)
    add_run(tp1, "PROPOSTA COMERCIAL", bold=True, color=NAVY, size=20, font='Calibri')

    tp2 = title_cell.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp2.paragraph_format.space_after = Pt(0)
    subtitle = f"{'Produção audiovisual' if not event else event} — {client}" if client else "Produção audiovisual"
    add_run(tp2, subtitle, color=GRAY, size=11, font='Calibri')

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── TABELA DE DADOS ───────────────────────────────────────────────────
    info_rows = [("Contratante", client or "—")]
    if cnpj:     info_rows.append(("CNPJ", cnpj))
    if contact:  info_rows.append(("Contato responsável", contact))
    if date:     info_rows.append(("Data provável de captação", date))
    if delivery: info_rows.append(("Entrega final", delivery))
    info_rows.append(("Investimento total", brl(total)))

    info_tbl = doc.add_table(rows=len(info_rows), cols=2)
    info_tbl.style = 'Table Grid'
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    COL_W = [Cm(5.5), Cm(11)]
    border_clr = 'CCCCCC'

    for i, (lbl, val) in enumerate(info_rows):
        cl = info_tbl.rows[i].cells[0]
        cr = info_tbl.rows[i].cells[1]
        cl.width = Cm(5.5)
        cr.width = Cm(11)
        set_cell_bg(cl, 'F5F5F5')
        cell_padding(cl, 60, 60, 100, 100)
        cell_padding(cr, 60, 60, 100, 100)
        set_cell_borders(cl, color=border_clr, size=4)
        set_cell_borders(cr, color=border_clr, size=4)

        pl = cl.paragraphs[0]
        pl.paragraph_format.space_before = Pt(0)
        pl.paragraph_format.space_after  = Pt(0)
        add_run(pl, lbl, color=GRAY, size=10.5, font='Calibri')

        pr = cr.paragraphs[0]
        pr.paragraph_format.space_before = Pt(0)
        pr.paragraph_format.space_after  = Pt(0)
        if i == len(info_rows)-1:
            add_run(pr, val, bold=True, color=NAVY, size=11, font='Calibri')
        else:
            add_run(pr, val, color=BLACK, size=10.5, font='Calibri')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── CAIXA OBJETO ──────────────────────────────────────────────────────
    obj_tbl = doc.add_table(rows=1, cols=1)
    obj_tbl.style = 'Table Grid'
    obj_cell = obj_tbl.rows[0].cells[0]
    obj_cell.width = Cm(16.5)
    set_cell_bg(obj_cell, 'EFF8FC')
    cell_padding(obj_cell, 80, 80, 120, 120)
    # Cyan border
    tc = obj_cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '0090BB')
        tcBorders.append(el)
    tcPr.append(tcBorders)

    op = obj_cell.paragraphs[0]
    op.paragraph_format.space_before = Pt(0)
    op.paragraph_format.space_after  = Pt(4)
    add_run(op, "Objeto", bold=True, color=CYAN, size=10.5, font='Calibri')

    op2 = obj_cell.add_paragraph()
    op2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    op2.paragraph_format.space_before = Pt(0)
    op2.paragraph_format.space_after  = Pt(0)
    add_run(op2, objeto, color=BLACK, size=10.5, font='Calibri')

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── 1. ESCOPO ─────────────────────────────────────────────────────────
    section_title(doc, 1, "ESCOPO DOS SERVIÇOS")
    if escopo:
        for item in escopo:
            bullet_item(doc, item)
    else:
        body_para(doc, f"A presente proposta contempla pré-produção, captação audiovisual e pós-produção, conforme os itens abaixo.")
        for t in team_items:
            bullet_item(doc, f"{t['qty']}x {t['label']} — {diarias} diária{'s' if diarias>1 else ''}")
        for e in edicao_items:
            bullet_item(doc, f"Pós-produção: {e['label']}")

    # ── 2. ENTREGÁVEIS ────────────────────────────────────────────────────
    section_title(doc, 2, "ENTREGÁVEIS")
    if entregaveis:
        for item in entregaveis:
            bullet_item(doc, item)
    else:
        for e in edicao_items:
            bullet_item(doc, f"Edição: {e['label']}")
        if not edicao_items:
            bullet_item(doc, "Material editado e finalizado conforme escopo acordado.")

    # ── 3. INVESTIMENTO ───────────────────────────────────────────────────
    section_title(doc, 3, "INVESTIMENTO E CONDIÇÕES DE PAGAMENTO")
    body_para(doc, f"O valor total desta proposta é de {brl(total)}, conforme detalhamento abaixo.")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Breakdown
    bd_rows = []
    for t in team_items:
        bd_rows.append((f"{t['qty']}x {t['label']} × {diarias} diária{'s' if diarias>1 else ''}", brl(t.get('price_total',0))))
    if transp > 0:
        bd_rows.append(("Transporte + alimentação", brl(transp)))
    for e in edicao_items:
        bd_rows.append((f"Edição: {e['label']}", brl(e.get('price',0))))
    bd_rows.append(("Subtotal", brl(subtotal)))
    bd_rows.append(("Caixa empresa (10%)", brl(caixa)))
    bd_rows.append(("NF / Impostos (6%)", brl(imposto)))
    bd_rows.append(("TOTAL", brl(total)))

    bd_tbl = doc.add_table(rows=len(bd_rows), cols=2)
    bd_tbl.style = 'Table Grid'
    for i,(lbl,val) in enumerate(bd_rows):
        cl = bd_tbl.rows[i].cells[0]
        cr = bd_tbl.rows[i].cells[1]
        cl.width = Cm(12); cr.width = Cm(4.5)
        cell_padding(cl,50,50,100,100); cell_padding(cr,50,50,100,100)
        set_cell_borders(cl, color=border_clr, size=4)
        set_cell_borders(cr, color=border_clr, size=4)
        is_total = (lbl == "TOTAL")
        is_sub   = (lbl == "Subtotal")
        is_muted = "%" in lbl
        if is_total:
            set_cell_bg(cl,'EFF8FC'); set_cell_bg(cr,'EFF8FC')
        pl = cl.paragraphs[0]; pl.paragraph_format.space_before=Pt(0); pl.paragraph_format.space_after=Pt(0)
        pr = cr.paragraphs[0]; pr.paragraph_format.space_before=Pt(0); pr.paragraph_format.space_after=Pt(0)
        pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        txt_color = GRAY if is_muted else (NAVY if is_total else BLACK)
        add_run(pl, lbl, bold=is_total or is_sub, color=txt_color, size=10.5 if not is_total else 11)
        add_run(pr, val, bold=is_total or is_sub, color=txt_color, size=10.5 if not is_total else 11)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Parcelas
    if parcelas:
        p_lbl = doc.add_paragraph()
        p_lbl.paragraph_format.space_after = Pt(4)
        add_run(p_lbl, "Cronograma de pagamento:", bold=True, color=BLACK, size=10.5)

        parc_tbl = doc.add_table(rows=len(parcelas)+1, cols=3)
        parc_tbl.style = 'Table Grid'
        # Header row
        hdrs = ["Parcela","Valor","Vencimento"]
        widths = [Cm(3.5), Cm(4), Cm(9)]
        for j,h in enumerate(hdrs):
            c = parc_tbl.rows[0].cells[j]; c.width=widths[j]
            set_cell_bg(c,'1A1A2E')
            cell_padding(c,50,50,100,100)
            set_cell_borders(c,color='1A1A2E',size=4)
            pp = c.paragraphs[0]; pp.paragraph_format.space_before=Pt(0); pp.paragraph_format.space_after=Pt(0)
            add_run(pp, h, bold=True, color=WHITE, size=10, font='Calibri')
        # Data rows
        for i,parc in enumerate(parcelas, 1):
            vals = [parc.get('num',''), brl(parc.get('valor',0)), parc.get('vencimento','')]
            for j,v in enumerate(vals):
                c = parc_tbl.rows[i].cells[j]; c.width=widths[j]
                if i%2==0: set_cell_bg(c,'F8F8F8')
                cell_padding(c,50,50,100,100)
                set_cell_borders(c,color=border_clr,size=4)
                pp = c.paragraphs[0]; pp.paragraph_format.space_before=Pt(0); pp.paragraph_format.space_after=Pt(0)
                add_run(pp, v, color=BLACK, size=10)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    body_para(doc, "Os pagamentos deverão ser realizados por meio previamente acordado entre as partes. Em caso de atraso, incidirão multa de 2% sobre o valor devido e juros de mora de 1% ao mês, calculados proporcionalmente até a regularização.")

    # ── 4. PRAZO ──────────────────────────────────────────────────────────
    section_title(doc, 4, "PRAZO DE ENTREGA")
    body_para(doc, f"A entrega final está prevista para {delivery or 'data a confirmar'}, condicionada ao envio do acervo, ao cumprimento do cronograma de produção e ao retorno das aprovações dentro dos prazos necessários para continuidade da edição.")
    body_para(doc, "Havendo atraso no envio de materiais, mudanças de escopo, demora em feedbacks ou atraso nas aprovações por parte da contratante, o prazo de entrega será automaticamente reajustado na mesma proporção do impacto causado.")

    # ── 5. RODADAS ────────────────────────────────────────────────────────
    section_title(doc, 5, "RODADAS DE AJUSTE")
    body_para(doc, "Estão incluídas 3 etapas de ajuste, compreendidas como:")
    for item in ["Primeiro corte.", "Ajustes iniciais.", "Ajustes finais."]:
        bullet_item(doc, item)
    body_para(doc, "As rodadas contemplam correções, refinamentos e alterações pontuais sobre a estrutura apresentada. Não incluem refação integral da montagem, mudança substancial de roteiro ou conceito após o início da edição. Demandas adicionais poderão ser orçadas à parte.")

    # ── 6. OBRIGAÇÕES ─────────────────────────────────────────────────────
    section_title(doc, 6, "OBRIGAÇÕES DA CONTRATANTE")
    for item in [
        "Disponibilizar o acervo histórico necessário para o projeto.",
        "Garantir que possui autorização para uso de imagens, marcas, depoimentos e demais conteúdos fornecidos.",
        "Centralizar aprovações, preferencialmente por meio de um responsável definido.",
        "Repassar informações, ajustes e validações dentro dos prazos acordados.",
    ]:
        bullet_item(doc, item)

    # ── 7. CAPTAÇÃO ───────────────────────────────────────────────────────
    section_title(doc, 7, "CAPTAÇÃO, LOCAÇÃO E DESLOCAMENTO")
    loc_txt = f"A presente proposta considera {diarias} diária{'s' if diarias>1 else ''} de captação"
    if location: loc_txt += f" em {location}"
    loc_txt += ", conforme alinhamento prévio. Qualquer necessidade de ampliação de tempo, nova diária ou deslocamentos adicionais não previstos poderá ser cobrada à parte, mediante aprovação prévia."
    body_para(doc, loc_txt)

    # ── 8. FORMATOS ───────────────────────────────────────────────────────
    section_title(doc, 8, "FORMATOS, ENTREGA E USO")
    for item in [
        "Os materiais serão entregues digitalmente, em formato compatível com o uso acordado.",
        "Os arquivos brutos não estão incluídos neste orçamento.",
        "A contratada poderá utilizar trechos, frames e versões do material para portfólio, divulgação profissional e apresentação comercial, salvo manifestação expressa em contrário por escrito antes da publicação.",
    ]:
        bullet_item(doc, item)

    # ── 9. RESPONSABILIDADE ───────────────────────────────────────────────
    section_title(doc, 9, "RESPONSABILIDADE SOBRE CONTEÚDOS E AUTORIZAÇÕES")
    body_para(doc, "A contratante declara ser responsável pelas autorizações de uso de imagem, voz, marcas, depoimentos, acervo histórico e demais materiais de terceiros fornecidos para edição. A contratada não se responsabiliza por eventual uso indevido de conteúdo encaminhado sem as devidas autorizações.")

    # ── 10. CANCELAMENTO ──────────────────────────────────────────────────
    section_title(doc, 10, "CANCELAMENTO, SUSPENSÃO E REMARCAÇÃO")
    body_para(doc, "Em caso de cancelamento do projeto após o aceite e início da pré-produção, os valores já pagos poderão ser retidos proporcionalmente ao estágio executado do trabalho. Em caso de suspensão ou paralisação por iniciativa da contratante após o início dos trabalhos, a contratada poderá faturar proporcionalmente as etapas já realizadas. Eventual remarcação de captação dependerá de disponibilidade de agenda.")

    if notes:
        section_title(doc, 11, "OBSERVAÇÕES")
        body_para(doc, notes)

    # ── ACEITE ────────────────────────────────────────────────────────────
    # Linha separadora
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_before = Pt(16)
    p_hr.paragraph_format.space_after  = Pt(4)
    pPr = p_hr._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:top')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'0090BB')
    pBdr.append(bot); pPr.append(pBdr)

    section_title(doc, 12, "ACEITE")
    body_para(doc, "O aceite desta proposta poderá ocorrer por assinatura física, assinatura digital, resposta formal por e-mail ou confirmação expressa por WhatsApp, implicando concordância integral com os termos aqui descritos.")

    # Espaço para assinaturas
    for _ in range(3):
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Tabela de assinatura
    sig_tbl = doc.add_table(rows=3, cols=3)
    sig_tbl.style = 'Table Grid'
    # Remove all borders
    for row in sig_tbl.rows:
        for cell in row.cells:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
            tcB = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right','insideH','insideV']:
                el = OxmlElement(f'w:{side}'); el.set(qn('w:val'),'none'); tcB.append(el)
            tcPr.append(tcB)

    sig_tbl.columns[0].width = Cm(7)
    sig_tbl.columns[1].width = Cm(2.5)
    sig_tbl.columns[2].width = Cm(7)

    # Row 0: signature lines
    for j in (0, 2):
        c = sig_tbl.rows[0].cells[j]
        p = c.paragraphs[0]; p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'),'single'); top.set(qn('w:sz'),'6')
        top.set(qn('w:space'),'1'); top.set(qn('w:color'),'333333')
        pBdr.append(top); pPr.append(pBdr)

    # Row 1: names
    names = [
        ("VIVAI STUDIO / Mateus de Oliveira", "Contratada"),
        ("", ""),
        (client or "Contratante", "Contratante"),
    ]
    for j,(name,role) in enumerate(names):
        c = sig_tbl.rows[1].cells[j]
        if name:
            p = c.paragraphs[0]; p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(0)
            add_run(p, name, bold=True, color=BLACK, size=10)
            p2 = c.add_paragraph(); p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(0)
            add_run(p2, role, color=GRAY, size=9)

    # Row 2: CNPJ/contact info
    if cnpj or contact:
        info_txt = f"Contato da contratante: {contact}" if contact else ""
        if cnpj: info_txt += f"  •  CNPJ: {cnpj}"
        c_info = sig_tbl.rows[2].cells[2]
        p_i = c_info.paragraphs[0]; p_i.paragraph_format.space_before=Pt(2); p_i.paragraph_format.space_after=Pt(0)
        add_run(p_i, info_txt.strip(), color=GRAY, size=8.5)

    # Rodapé do documento
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_end, f"Proposta gerada em {datetime.now().strftime('%d/%m/%Y')}  •  VIVAI Studio  •  www.studiovivai.com",
            color=GRAY, size=8.5)

    doc.save(output_path)
    print(f"✅ DOCX gerado: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Uso: python3 generate_docx.py <json_file_ou_json> output.docx [logo.png]")
        sys.exit(1)
    json_arg = sys.argv[1]
    output   = sys.argv[2]
    logo     = sys.argv[3] if len(sys.argv) > 3 else None
    if os.path.isfile(json_arg):
        with open(json_arg,"r",encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = json.loads(json_arg)
    generate(data, output, logo)
